from tkinter import Text
from tkinter import ttk
from googletrans import Translator  # Make sure to install the 'googletrans==4.0.0-rc1' library
from tkinter import RIDGE
import tkinter as tk
from tkinter import messagebox
import docx
import os
import subprocess
import platform

class MainMenu:
    def __init__(self, root):
        self.root = root
        self.root.title("Main Menu")

        self.root.geometry("750x550")

        self.root.configure(bg="#001F3F")

        icon_image = tk.PhotoImage(file="LanguageHUB_icon.png") 
        resized_icon = icon_image.subsample(6, 6)

        icon_label = tk.Label(root, image=resized_icon, bg="#001F3F")
        icon_label.image = resized_icon
        icon_label.pack(pady=10)

        welcome_label = tk.Label(root, text="LanguageHUB", font=("DIN Alternate", 72, "bold"), fg="white", bg="#001F3F")
        welcome_label.pack(pady=0)

        title_button_style = {"font": ("Arial", 16, "bold"), "fg": "#003366", "bg": "#001F3F", "relief": tk.GROOVE, "bd": 3, "padx": 15, "pady": 10}
        close_button_style = {"font": ("Arial", 12, "bold"), "fg": "#003366", "bg": "#001F3F", "relief": tk.GROOVE, "bd": 3, "padx": 10, "pady": 5}

        text_translation_button = tk.Button(root, text="Text Translation", command=self.text_translation, **title_button_style)
        text_translation_button.pack(pady=20)

        document_translation_button = tk.Button(root, text="Document Translation", command=self.document_translation, **title_button_style)
        document_translation_button.pack(pady=20)

        recording_translation_button = tk.Button(root, text="Recording Translation", command=self.record_translation, **title_button_style)
        recording_translation_button.pack(pady=20)

        close_button = tk.Button(root, text="Close", command=self.close_window, **close_button_style)
        close_button.pack(side=tk.LEFT, anchor=tk.SW, padx=10, pady=10)

    def text_translation(self):
        text_translation_window = tk.Toplevel(self.root)
        yey = texttranslator_window(text_translation_window)

    def document_translation(self):
        doc_translation_window = tk.Toplevel(self.root)
        yey = doctranslator_window(doc_translation_window)
    def record_translation(self):
        messagebox.showinfo("Recording Translation", "This feature will be coming soon. Stay tuned!")
    def close_window(self):
        if messagebox.askokcancel("Confirmation", "Are you sure you want to close LanguageHUB?"):
            self.root.destroy()

class texttranslator_window:
    def __init__(self, root):
        self.root = root
        self.root.title("Translator")
        self.root.geometry("750x550")
        self.root.resizable(False, False)
        self.root.configure(bg="#001F3F")

        frame = tk.Frame(self.root, bg="#001F3F")
        frame.place(relx=0.5, rely=0, anchor="n")
        label = tk.Label(frame, text="LanguageHUB", font=('DIN Alternate', 42, 'bold'), pady = 10, fg="white", bg="#001F3F")
        label.pack()

        icon_image1 = tk.PhotoImage(file='LanguageHUB_icon.png')
        resized_image1 = icon_image1.subsample(7, 7)
        icon_label1 = tk.Label(self.root, image=resized_image1, bg="#001F3F")
        icon_label1.image = resized_image1
        icon_label1.place(relx=1, rely=0.01, anchor="ne")

        icon_image2 = tk.PhotoImage(file='LanguageHUB_icon.png')
        resized_image2 = icon_image2.subsample(7, 7)
        icon_label2 = tk.Label(self.root, image=resized_image2, bg="#001F3F")
        icon_label2.image = resized_image2
        icon_label2.place(relx=0.11, rely=0.01, anchor="ne")

        white_frame = tk.Frame(self.root, bg="white", height=3, width=750)
        white_frame.place(relx=0, rely=0.17, anchor="nw")

        close_button_style = {"font": ("DIN Alternate", 12, "bold"), "fg": "#003366", "bg": "#001F3F", "relief": tk.GROOVE, "bd": 3, "padx": 10, "pady": 5}
        close_button = tk.Button(self.root, text="Back", command=self.root.destroy, **close_button_style)
        close_button.place(x=10, y=500)

        Cb_autoDetect = ttk.Combobox(self.root, width=20, font=('DIN Alternate', 11), state="readonly")
        Cb_autoDetect['values'] = ('Auto Detect',)
        Cb_autoDetect.place(x=100, y=120)
        Cb_autoDetect.current(0)

        self.cb_choose_langauge = ttk.Combobox(self.root, width=20, font=('DIN Alternate', 11), state="readonly")
        self.cb_choose_langauge['values'] = (
            'Afrikaans', 'Albanian', 'Arabic', 'Armenian', 'Basque', 'Belarusian', 'Bengali', 'Bosnian', 'Bulgarian',
            'Cebuano', 'Chichewa', 'Corsican', 'Croatian', 'Danish', 'Dutch', 'English', 'Esperanto', 'Estonian',
            'Filipino', 'Finnish', 'French', 'Frisian', 'Galician', 'Georgian', 'German', 'Greek', 'Gujarati', 'Haitian Creole',
            'Hausa', 'Hawaiian', 'Hebrew', 'Hindi', 'Hmong', 'Hungarian', 'Icelandic', 'Igbo', 'Indonesian', 'Irish', 'Italian',
            'Japanese', 'Javanese', 'Kannada', 'Kazakh', 'Khmer', 'Kinyarwanda', 'Korean', 'Kurdish', 'Kyrgyz', 'Lao', 'Latin',
            'Latvian', 'Lithuanian', 'Luxembourgish', 'Macedonian', 'Malagasy', 'Malay', 'Malayalam', 'Maltese', 'Maori',
            'Marathi', 'Mongolian', 'Myanmar', 'Nepali', 'Norwegian', 'Odia', 'Pashto', 'Persian', 'Polish', 'Portuguese',
            'Punjabi', 'Romanian', 'Russian', 'Samoan', 'Scots Gaelic', 'Serbian', 'Sesotho', 'Shona', 'Sindhi', 'Sinhala',
            'Slovak', 'Slovenian', 'Somali', 'Spanish', 'Sundanese', 'Swahili', 'Swedish', 'Tajik', 'Tamil', 'Tatar', 'Telugu',
            'Thai', 'Turkish', 'Turkmen', 'Ukrainian', 'Urdu', 'Uyghur', 'Uzbek', 'Vietnamese', 'Welsh', 'Xhosa', 'Yiddish',
            'Yoruba', 'Zulu',
        )
        self.cb_choose_langauge.place(x=500, y=120)
        self.cb_choose_langauge.current(0)

        self.textbox1 = Text(root, width=45, height=10, borderwidth=4, font=('DIN Alternate', 14), pady=5, relief=RIDGE)
        self.textbox1.place(x=20, y=170)
        self.textbox2 = Text(self.root, width=45, height=10, borderwidth=4, font=('DIN Alternate', 14), pady=5, relief=RIDGE)
        self.textbox2.place(x=400, y=170)

        translate_button = tk.Button(self.root, text="Translate", borderwidth=2, **close_button_style, cursor="hand2", command=self.translate, width=8)
        translate_button.place(x=240, y=410)

        clear_button = tk.Button(self.root, text="Clear", **close_button_style, borderwidth=2,
                             cursor="hand2", command=self.clear, width=8)
        clear_button.place(x=400, y=410)

    def translate(self):
        txt_input = self.textbox1.get("1.0", "end-1c")
        language = self.cb_choose_langauge.get()

        if txt_input == '':
            messagebox.showerror('Translator - error', 'No text. Please fill the box')
        else:
            self.textbox2.delete(1.0, 'end')
            translator = Translator()
            output = translator.translate(txt_input, dest=language)
            self.textbox2.insert('end', output.text)

    def clear(self):
        self.textbox1.delete(1.0, 'end')
        self.textbox2.delete(1.0, 'end')

class doctranslator_window:
    def __init__(self, root):
        self.root = root
        self.root.title("Translator")
        self.root.geometry("750x550")
        self.root.resizable(False, False)
        self.root.configure(bg="#001F3F")
        
        frame = tk.Frame(self.root, bg="#001F3F")
        frame.place(relx=0.5, rely=0.01, anchor="n")
        label = tk.Label(frame, text="LanguageHUB", font=('DIN Alternate', 42, 'bold'), pady = 10, fg="white", bg="#001F3F")
        label.pack()

        icon_image1 = tk.PhotoImage(file='LanguageHUB_icon.png')
        resized_image1 = icon_image1.subsample(7, 7)
        icon_label1 = tk.Label(self.root, image=resized_image1, bg="#001F3F")
        icon_label1.image = resized_image1
        icon_label1.place(relx=1, rely=0.01, anchor="ne")

        icon_image2 = tk.PhotoImage(file='LanguageHUB_icon.png')
        resized_image2 = icon_image2.subsample(7, 7)
        icon_label2 = tk.Label(self.root, image=resized_image2, bg="#001F3F")
        icon_label2.image = resized_image2
        icon_label2.place(relx=0.11, rely=0.01, anchor="ne")

        white_frame = tk.Frame(self.root, bg="white", height=3, width=750)
        white_frame.place(relx=0, rely=0.17, anchor="nw")

        frame = tk.Frame(self.root, bg="#001F3F")
        frame.place(relx=0.265, rely=0.3, anchor="n")
        label = tk.Label(frame, text="Name of the file you want to translate: ", font=('Kohinur Bangla', 20), pady = 10, fg="white", bg="#001F3F")
        label.pack()

        frame = tk.Frame(self.root, bg="#001F3F")
        frame.place(relx=0.18, rely=0.43, anchor="n")
        label = tk.Label(frame, text="Language to translate:  ", font=('Kohinur Bangla', 20), pady = 10, fg="white", bg="#001F3F")
        label.pack()

        frame = tk.Frame(self.root, bg="#001F3F")
        frame.place(relx=0.2, rely=0.56, anchor="n")
        label = tk.Label(frame, text="How do you want to call the\ntranslated doc?: ",justify="left", font=('Kohinur Bangla', 20), pady = 10, fg="white", bg="#001F3F")
        label.pack()

        other_button_style = {"font": ("DIN Alternate", 18, "bold"), "fg": "#003366", "bg": "#001F3F", "relief": tk.GROOVE, "bd": 3, "padx": 10, "pady": 5}

        close_button_style = {"font": ("DIN Alternate", 12, "bold"), "fg": "#003366", "bg": "#001F3F", "relief": tk.GROOVE, "bd": 3, "padx": 10, "pady": 5}
        close_button = tk.Button(self.root, text="Back", command=self.root.destroy, **close_button_style)
        close_button.place(x=10, y=500)

        self.cb_choose_langauge = ttk.Combobox(self.root, width=30, height = 5, font=('DIN Alternate', 14), state="readonly")
        self.cb_choose_langauge['values'] = (
            'Afrikaans', 'Albanian', 'Arabic', 'Armenian', 'Basque', 'Belarusian', 'Bengali', 'Bosnian', 'Bulgarian',
            'Cebuano', 'Chichewa', 'Corsican', 'Croatian', 'Danish', 'Dutch', 'English', 'Esperanto', 'Estonian',
            'Filipino', 'Finnish', 'French', 'Frisian', 'Galician', 'Georgian', 'German', 'Greek', 'Gujarati', 'Haitian Creole',
            'Hausa', 'Hawaiian', 'Hebrew', 'Hindi', 'Hmong', 'Hungarian', 'Icelandic', 'Igbo', 'Indonesian', 'Irish', 'Italian',
            'Japanese', 'Javanese', 'Kannada', 'Kazakh', 'Khmer', 'Kinyarwanda', 'Korean', 'Kurdish', 'Kyrgyz', 'Lao', 'Latin',
            'Latvian', 'Lithuanian', 'Luxembourgish', 'Macedonian', 'Malagasy', 'Malay', 'Malayalam', 'Maltese', 'Maori',
            'Marathi', 'Mongolian', 'Myanmar', 'Nepali', 'Norwegian', 'Odia', 'Pashto', 'Persian', 'Polish', 'Portuguese',
            'Punjabi', 'Romanian', 'Russian', 'Samoan', 'Scots Gaelic', 'Serbian', 'Sesotho', 'Shona', 'Sindhi', 'Sinhala',
            'Slovak', 'Slovenian', 'Somali', 'Spanish', 'Sundanese', 'Swahili', 'Swedish', 'Tajik', 'Tamil', 'Tatar', 'Telugu',
            'Thai', 'Turkish', 'Turkmen', 'Ukrainian', 'Urdu', 'Uyghur', 'Uzbek', 'Vietnamese', 'Welsh', 'Xhosa', 'Yiddish',
            'Yoruba', 'Zulu',
        )
        self.cb_choose_langauge.place(x=400, y=260)
        self.cb_choose_langauge.current(0)

        self.textbox1 = Text(self.root, width=30, height=0.8, borderwidth=4, font=('DIN Alternate', 14), pady=5, relief=RIDGE)
        self.textbox1.place(x=405, y=170)

        self.textbox2 = Text(self.root, width=30, height=0.8, borderwidth=4, font=('DIN Alternate', 14), pady=5, relief=RIDGE)
        self.textbox2.place(x=405, y=330)

        translate_button = tk.Button(self.root, text="Translate", borderwidth=2, **other_button_style, cursor="hand2", command=self.translate2, width=8)
        translate_button.place(x=240, y=430)

        clear_button = tk.Button(self.root, text="Clear All", **other_button_style, borderwidth=2,
                             cursor="hand2", command=self.clear, width=8)
        clear_button.place(x=400, y=430)


    def translate2(self):
        sourcename = self.textbox1.get("1.0", "end-1c")
        targetlan = self.cb_choose_langauge.get()
        targetdocname = self.textbox2.get("1.0", "end-1c")
        doc = docx.Document(sourcename + ".docx")
        translated_doc = docx.Document()
        paragraphs = [para.text for para in doc.paragraphs]
        translator = Translator()

        if sourcename == '':
            messagebox.showerror('Translator - error', 'Please fill the box of the document name')
        else:
            for i, para in enumerate(paragraphs):
                try:  # We used try to avoid any exception errors
                    translation = translator.translate(para,
                                                       dest=targetlan)  # Calling the translator function that will take it element/parragraph in the list parragraphs and traduce it to the targetlan
                    translated_doc.add_paragraph(
                        translation.text)  # Adds the translated parragraphs to the translated_doc previously created
                    messagebox.showinfo(
                        "Translation - Completed","Document translation is completed. Once again, LanguageHUB has saved you time.")  # A confirmation message
                except:  # In case of error; such as a not uploaded document or not readable document such as a .pwt
                    messagebox.showerror(
                        'Something went wronf',"The program did not work. Make sure the document is uploaded in the environment as a .docx and that the name of the document matches to the one it was input when the system asked for it")  # Message telling that the program did not work and reminding the user that everything has to be uploaded as a .docx file and the name must match when the system asks for it
            translated_doc.save(
                targetdocname + ".docx")
            current_directory = os.getcwd()
            file_directory = current_directory + '/' + targetdocname + ".docx"
            def open_word_document(docx_file_path):
                system_platform = platform.system()

                if system_platform == "Windows":
                    # Windows
                    subprocess.Popen(['start', '""', docx_file_path], shell=True)
                elif system_platform == "Darwin":
                    # macOS
                    subprocess.Popen(['open', docx_file_path])
                elif system_platform == "Linux":
                    # Linux
                    subprocess.Popen(['xdg-open', docx_file_path])

            open_word_document(file_directory)

    def clear(self):
        self.textbox1.delete(1.0, 'end')
        self.textbox2.delete(1.0, 'end')

if __name__ == "__main__":
    root = tk.Tk()
    app = MainMenu(root)
    root.mainloop()


